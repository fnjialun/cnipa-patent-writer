# -*- coding: utf-8 -*-
"""build_patent.py — 套用模板格式装配中国发明专利 docx。

做法:克隆模板里真实段落作"原型"(版式 100% 继承)+ 复刻四节四运行页眉(说明书摘要/权利要求书/说明书/
说明书附图)+ 逐节复刻模板真实的页脚/页码 + 嵌入配图 + 清掉克隆带入的孤儿图 + **剥离模板内嵌字体子集**。
**只借格式、不抄内容**——文字由调用方据本技术方案提供。

内嵌字体的坑(跨主题套模板必踩):模板若用 Word"嵌入字体"只存了模板原有文字的字形子集,克隆后本专利的新文字
不在子集内会变豆腐块(尤其黑体发明名称)。save() 已自动剥离内嵌、改用系统完整字体(见 _strip_embedded_fonts)。

页码(开局从模板逐节扒真实配置,save() 时原样复刻,不做硬编码假设):模板某节挂 PAGE 域页脚就挂它、
有 pgNumType 就照其 start 重起、是空页脚就挂空页脚(该节无页码)、无 footerReference 就不挂(继承前节页脚续接)。
不同模板对摘要/附图页码的处理本就不同(有的附图独立无页码、有的续接说明书页码),本脚本据 inspect 出的实测值
一一还原——模板怎样,产出就怎样。完稿务必渲染逐页核对页码(见 render_check.py)。

依赖: pip install python-docx

用法:
    from build_patent import PatentBuilder
    b = PatentBuilder("template.docx")
    b.abstract("本发明公开了……")                  # 摘要(可多段) -> 说明书摘要节
    for c in claim_paragraphs: b.claim(c)          # 每条权利要求段 -> 权利要求书节
    b.spec_title("一种……方法及系统")              # 说明书标题(技术领域前,居中) -> 以下为说明书节
    b.heading("技术领域"); b.body_justify("本发明涉及……")
    b.heading("背景技术"); b.body_justify("……")
    b.heading("发明内容"); b.body("……")
    b.body("有益效果："); b.body("本发明……")
    b.heading("附图说明"); b.body("图1为……示意图；")
    b.heading("具体实施方式"); b.body("下面结合……"); b.body("为了让……")
    b.subhead("实施例1："); b.body("步骤S1：……"); b.body("在本步骤中，……")
    b.figure("fig1.png", 4.8); b.caption("图1")    # -> 说明书附图节
    b.save("输出.docx")

注意:必须按 abstract* -> claim* -> spec_title/heading/body/subhead* -> figure/caption* 的顺序调用,
节的归属据此判定。各方法返回所建段落元素(一般无需用)。
"""
import os
import re
import copy
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEADINGS = ("技术领域", "背景技术", "发明内容", "具体实施方式", "附图说明")
HEADER_PARTS = ("说明书摘要", "权利要求书", "说明书", "说明书附图")


def _norm(s):
    return (s or "").replace(" ", "").replace("　", "").strip()


class PatentBuilder:
    def __init__(self, template_path):
        self.doc = Document(template_path)
        self._body = self.doc.element.body
        sp = self._body.find(qn("w:sectPr"))
        self.base_sectPr = copy.deepcopy(sp)
        self.hdr_rid = self._header_rids()
        self.sect_cfg = self._capture_section_configs()  # 须在清空 body 前抓(此时各节 sectPr 还在)
        self.proto = self._capture_prototypes()
        # 清空正文(暂移除 body 级 sectPr,保存时按四节重建)
        for el in list(self._body):
            if el is not sp:
                self._body.remove(el)
        if sp is not None:
            self._body.remove(sp)
        self._abs_last = None
        self._claims_last = None
        self._spec_last = None
        self._warn = []

    # ---------- 模板探查 ----------
    def _header_rids(self):
        d = {}
        for rid, rel in self.doc.part.rels.items():
            if rel.reltype.endswith("/header"):
                txt = _norm("".join((x.text or "") for x in rel.target_part.element.iter(qn("w:t"))))
                if txt:
                    d[txt] = rid
        return d

    def _capture_section_configs(self):
        """逐节扒模板真实的"页脚引用 + 页码起始",按各节页眉文字归档,供 _make_sectPr **原样复刻**。
        这是真正的"版式继承"——模板某节怎样处理页脚/页码,产出就怎样,不做任何假设。
        返回 {归一页眉文字: {"ftr_rid": 该节 footerReference 的 rId 或 None, "pgnum_start": "1" 或 None}}。
        注:模板里"带 PAGE 域的页脚"会出页码;"空页脚"会抑制"链接到前一节"使该节无页码;
        "无 footerReference"则继承前一节页脚(常见于附图节续接说明书页码)。三种情形本方法都如实记录。"""
        rid2txt = {rid: txt for txt, rid in self.hdr_rid.items()}
        cfg = {}
        for sp in self._body.iter(qn("w:sectPr")):
            hr = sp.find(qn("w:headerReference"))
            key = rid2txt.get(hr.get(qn("r:id"))) if hr is not None else None
            if not key:
                continue
            fr = sp.find(qn("w:footerReference"))
            pg = sp.find(qn("w:pgNumType"))
            cfg[key] = {
                "ftr_rid": fr.get(qn("r:id")) if fr is not None else None,
                "pgnum_start": pg.get(qn("w:start")) if pg is not None else None,
            }
        return cfg

    def _capture_prototypes(self):
        """按中国专利标准结构识别 6 类段落原型(deepcopy 底层 w:p)。找不到则回退并告警。"""
        paras = self.doc.paragraphs
        proto = {}
        prev_nonempty = None
        for i, p in enumerate(paras):
            t = p.text.strip()
            if not t:
                continue
            if t in HEADINGS and "head" not in proto:
                proto["head"] = copy.deepcopy(p._p)
                if t == "技术领域":
                    # 标题 = 紧邻"技术领域"之前的非空段(居中发明名称)
                    if prev_nonempty is not None and "title" not in proto:
                        proto["title"] = copy.deepcopy(prev_nonempty._p)
                    # 缺省文本正文 = "技术领域"之后的下一非空段
                    for q in paras[i + 1:]:
                        if q.text.strip():
                            proto["bodyj"] = copy.deepcopy(q._p)
                            break
            if "body" not in proto and t.startswith("本发明"):
                proto["body"] = copy.deepcopy(p._p)
            if "sub" not in proto and re.match(r"^实施例\s*\d+", t):
                proto["sub"] = copy.deepcopy(p._p)
            if "cap" not in proto and re.match(r"^图\s*\d+$", t):
                proto["cap"] = copy.deepcopy(p._p)
            prev_nonempty = p
        # 回退:body 兜底其余
        if "body" not in proto:
            for p in paras:
                if p.text.strip():
                    proto["body"] = copy.deepcopy(p._p)
                    break
        for k in ("bodyj", "head", "title", "sub", "cap"):
            if k not in proto:
                proto[k] = copy.deepcopy(proto["body"])
                self._warn_later(f"未在模板中识别到原型[{k}],已回退用正文原型,请渲染核对。")
        return proto

    def _warn_later(self, msg):
        # __init__ 中 self._warn 还没建,延迟收集
        if not hasattr(self, "_warn"):
            self._warn = []
        self._warn.append(msg)

    # ---------- 段落克隆 ----------
    def _clone(self, key, text):
        p = copy.deepcopy(self.proto[key])
        pPr = p.find(qn("w:pPr"))
        rpr = None
        fr = p.find(qn("w:r"))
        if fr is not None:
            rp = fr.find(qn("w:rPr"))
            if rp is not None:
                rpr = copy.deepcopy(rp)
        for ch in list(p):
            if ch is not pPr:
                p.remove(ch)
        r = OxmlElement("w:r")
        if rpr is not None:
            r.append(rpr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        p.append(r)
        self._body.append(p)
        return p

    # ---------- 内容 API ----------
    def abstract(self, text):
        self._abs_last = self._clone("body", text)
        return self._abs_last

    def claim(self, text):
        self._claims_last = self._clone("body", text)
        return self._claims_last

    def spec_title(self, text):
        self._spec_last = self._clone("title", text)
        return self._spec_last

    def heading(self, text):
        self._spec_last = self._clone("head", text)
        return self._spec_last

    def body(self, text):
        self._spec_last = self._clone("body", text)
        return self._spec_last

    def body_justify(self, text):
        self._spec_last = self._clone("bodyj", text)
        return self._spec_last

    def subhead(self, text):
        self._spec_last = self._clone("sub", text)
        return self._spec_last

    def caption(self, text):
        return self._clone("cap", text)

    def figure(self, image_path, width_in):
        pic = self.doc.add_paragraph()
        pic.alignment = AL.CENTER
        pic.paragraph_format.space_before = Pt(6)
        pic.add_run().add_picture(image_path, width=Inches(width_in))
        return pic._p

    # ---------- 分节 ----------
    def _make_sectPr(self, header_text):
        """据 base_sectPr 复刻一节:换该节页眉、另起页,并**原样复刻模板该节真实的页脚引用与页码起始**。

        关键修复:base_sectPr 取自模板最后一节(附图节),若四节都照它复制会丢失正文节页码。
        故这里改为"按节查 sect_cfg(开局从模板扒的真实配置)逐节复刻"——模板某节有 PAGE 域页脚就挂它、
        有 pgNumType 就照其 start 重起、是空页脚就挂空页脚(抑制继承=无页码)、无页脚就不挂(继承前节续接)。
        不同模板对摘要/附图页码的处理本就不同(有的附图无页码独立、有的续接说明书),本方法据模板实测一一还原,
        不做硬编码假设。子元素顺序遵循 OOXML:headerRef → footerRef → type → pgSz → pgMar → pgNumType → cols。"""
        sp = copy.deepcopy(self.base_sectPr)
        for hr in sp.findall(qn("w:headerReference")):
            sp.remove(hr)
        for fr in sp.findall(qn("w:footerReference")):  # 清掉 base 带入的页脚/页码,改按本节 cfg 复刻
            sp.remove(fr)
        for pg in sp.findall(qn("w:pgNumType")):
            sp.remove(pg)
        key = _norm(header_text)
        rid = self.hdr_rid.get(key)
        last_anchor = None
        if rid:
            hr = OxmlElement("w:headerReference")
            hr.set(qn("w:type"), "default")
            hr.set(qn("r:id"), rid)
            sp.insert(0, hr)
            last_anchor = hr
        else:
            self._warn.append(f"模板缺运行页眉[{header_text}],该节沿用 base 页眉,请核对。")
        cfg = self.sect_cfg.get(key, {})
        if not cfg:
            self._warn.append(f"模板未扒到[{header_text}]节的页脚/页码配置,该节按无页脚处理,请核对。")
        # 复刻页脚引用(空页脚也要挂——它正是模板用来"该节无页码"的手段)
        if cfg.get("ftr_rid"):
            fr = OxmlElement("w:footerReference")
            fr.set(qn("w:type"), "default")
            fr.set(qn("r:id"), cfg["ftr_rid"])
            (last_anchor.addnext(fr) if last_anchor is not None else sp.insert(0, fr))
        # type=nextPage(排在 footerRef 之后、pgSz 之前)
        t = sp.find(qn("w:type"))
        if t is None:
            t = OxmlElement("w:type")
            pgSz = sp.find(qn("w:pgSz"))
            (pgSz.addprevious(t) if pgSz is not None else sp.append(t))
        t.set(qn("w:val"), "nextPage")
        # 复刻页码起始(模板该节有 pgNumType 才补,start 用模板真实值)
        if cfg.get("pgnum_start") is not None:
            pg = OxmlElement("w:pgNumType")
            pg.set(qn("w:start"), cfg["pgnum_start"])
            cols = sp.find(qn("w:cols"))
            grid = sp.find(qn("w:docGrid"))
            (cols.addprevious(pg) if cols is not None
             else grid.addprevious(pg) if grid is not None else sp.append(pg))
        return sp

    def _end_section(self, p_el, sect):
        if p_el is None:
            return
        pPr = p_el.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p_el.insert(0, pPr)
        pPr.append(sect)

    def _drop_orphan_images(self):
        used = set()
        for blip in self._body.iter(qn("a:blip")):
            e = blip.get(qn("r:embed"))
            if e:
                used.add(e)
        for rid, rel in list(self.doc.part.rels.items()):
            if "image" in rel.reltype and rid not in used:
                self.doc.part.drop_rel(rid)

    def _strip_embedded_fonts(self):
        """关掉并清除模板带入的"内嵌字体子集"。**这是跨主题套模板时最隐蔽的坑**:
        很多专利模板用 Word 的"将字体嵌入文件"功能,只内嵌了模板原有文字的字形子集(font*.odttf)。
        克隆后这些子集被沿用,而本专利的新文字(模板里没有的字)不在子集内——渲染/在 Word 打开时
        这些字会变成豆腐块□,尤以**发明名称等黑体行**最明显(其子集往往最小)。删掉内嵌、改用系统
        完整字体即可根治。做法:① settings 去掉 embedTrueTypeFonts/embedSystemFonts/saveSubsetFonts 开关;
        ② fontTable 各字体条目去掉 embedRegular/Bold/Italic/BoldItalic 引用;③ 丢弃 fonts/*.odttf 部件。"""
        # ① 功能性根治:settings 去掉内嵌开关——没有它,消费端就用系统完整字体,新文字不再豆腐块
        try:
            st = self.doc.settings.element
            for tag in ("w:embedTrueTypeFonts", "w:embedSystemFonts", "w:saveSubsetFonts"):
                for el in st.findall(qn(tag)):
                    st.remove(el)
        except Exception as e:
            self._warn.append(f"清内嵌字体开关失败({e}),请渲染核对发明名称等黑体字是否豆腐块。")
        # ② 顺手清理:删 fontTable 的 embed 引用并丢弃 .odttf 子集(减体积)。best-effort,失败不影响①
        try:
            from lxml import etree
            for rid, rel in list(self.doc.part.rels.items()):
                if rel.reltype.endswith("/fontTable"):
                    ft = rel.target_part
                    root = etree.fromstring(ft.blob)
                    for f in root.findall(qn("w:font")):
                        for ch in list(f):
                            if ch.tag.rsplit("}", 1)[-1] in (
                                    "embedRegular", "embedBold", "embedItalic", "embedBoldItalic"):
                                f.remove(ch)
                    ft._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                    for frid, frel in list(ft.rels.items()):
                        if frel.reltype.endswith("/font"):
                            ft.drop_rel(frid)
        except Exception as e:
            self._warn.append(f"清理内嵌字体子集部件失败({e}),不影响功能(内嵌开关已关)。")

    def save(self, out_path):
        # 四节:摘要/权利要求书/说明书 的 sectPr 挂到各节末段;附图节用 body 级 sectPr 收尾。
        # 每节的页脚/页码均按 sect_cfg 原样复刻模板对应节(见 _make_sectPr),不做硬编码假设。
        self._end_section(self._abs_last, self._make_sectPr("说明书摘要"))
        self._end_section(self._claims_last, self._make_sectPr("权利要求书"))
        self._end_section(self._spec_last, self._make_sectPr("说明书"))
        self._body.append(self._make_sectPr("说明书附图"))
        self._drop_orphan_images()
        self._strip_embedded_fonts()   # 关键:去掉模板内嵌的字体子集,否则新文字(尤其黑体发明名称)会豆腐块
        self.doc.save(out_path)
        if self._warn:
            print("[build_patent 告警]")
            for w in self._warn:
                print("  -", w)
        print("saved:", out_path, "| paragraphs:", len(self.doc.paragraphs))
        return out_path
