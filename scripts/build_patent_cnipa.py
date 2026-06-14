# -*- coding: utf-8 -*-
"""build_patent_cnipa.py — 无模板兜底装配器:区块 JSON → 中国发明专利 .docx(CNIPA 4 分节结构)。

何时用:**用户没有给模板 .docx** 时用本脚本(按 CNIPA 标准版式从零生成)。
若用户给了模板,优先用 build_patent.py(PatentBuilder,克隆模板段落,版式 100% 继承)。
两者产出同一套 4 分节结构,只是格式来源不同(本脚本=内置标准参数;PatentBuilder=模板实测)。

用法:
    python build_patent_cnipa.py blocks.json 输出文件名.docx

blocks.json 结构:
    {"title": "...", "blocks": [["kind","payload"], ...]}
  kind ∈ abstract | claim | title | h1 | h2 | body | code | figure
  figure 的 payload 可以是图片路径字符串,或 {"path": "..."} 的 JSON 串。
  **内容据本技术方案撰写,勿照搬任何模板/范文的技术内容。**

产出 4 分节:说明书摘要 / 权利要求书 / 说明书 / 说明书附图,
各节独立页眉(居中加粗15pt+下边框线)、权要与说明书各自页码从1起、各节另起页;
正文 宋体12pt·首行缩进2字·1.5倍行距;发明名称 黑体14pt居中;
章节标题14pt加粗左对齐段前后7.8pt;实施例X:/有益效果: 12pt加粗;
权利要求按";"分句成段、正文长段按句拆分、半角标点归全角。
版式参数说明见 references/cnipa-format-spec.md(以用户模板实测为准时改对应参数)。
"""
import sys, json, re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

SONG = "宋体"; HEI = "黑体"; MONO = "Courier New"
SEG_MIN = 150         # 正文长段拆分目标字数(按模板段长调,常见中位~68/最长~530)
MAX_FIG_W_CM = 15.3   # 附图最大宽(= A4 内容宽)
MAX_FIG_H_CM = 22.0   # 附图最大高(留页眉页脚)

# ---------------- 字体 / 段落基元 ----------------
def set_font(run, east=SONG, latin=SONG, size=12, bold=False):
    run.font.size = Pt(size); run.font.bold = bold; run.font.name = latin
    rf = run._element.get_or_add_rPr().get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), east); rf.set(qn("w:ascii"), latin); rf.set(qn("w:hAnsi"), latin)

def indent2(p):
    ind = p._p.get_or_add_pPr().get_or_add_ind()
    ind.set(qn("w:firstLineChars"), "200"); ind.set(qn("w:firstLine"), str(2 * 240))

def body_para(doc, text):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; indent2(p)
    set_font(p.add_run(text), size=12)

def heading(doc, text, kind):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if kind == "title":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(6); pf.space_after = Pt(10)
        set_font(p.add_run(text), east=HEI, latin=HEI, size=14, bold=True)
    elif kind == "h1":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(7.8); pf.space_after = Pt(7.8)
        set_font(p.add_run(text), size=14, bold=True)
    else:  # h2
        pf.space_before = Pt(0); pf.space_after = Pt(0); indent2(p)
        set_font(p.add_run(text), size=12, bold=True)

def add_code(doc, text):
    for ln in text.split("\n"):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(0); pf.space_after = Pt(0); pf.left_indent = Cm(0.6)
        set_font(p.add_run(ln if ln else " "), east=SONG, latin=MONO, size=9)

def add_figure(doc, path):
    w, h = Image.open(path).size; ratio = w / h
    w_cm = MAX_FIG_W_CM; h_cm = w_cm / ratio
    if h_cm > MAX_FIG_H_CM:
        h_cm = MAX_FIG_H_CM; w_cm = h_cm * ratio
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(path, width=Cm(w_cm))

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ---------------- 文本规整 ----------------
def norm_punct(s):
    """半角标点→全角(仅 CJK 上下文,避开 x1,y1 / 0.55 / v2.0 / http: 等技术记号)。"""
    C = r"一-鿿"
    s = re.sub(rf"(?<=[{C}]),", "，", s)
    s = re.sub(rf"(?<=[{C}]);", "；", s)
    s = re.sub(rf"(?<=[{C}]):(?=[{C}“\"])", "：", s)
    s = re.sub(rf"(?<=[{C}])\)", "）", s)
    s = re.sub(rf"\((?=[{C}])", "（", s)
    return s

def split_claim(text):
    """权利要求按 换行/";" 拆成短句段(对标模板逐句成段)。"""
    out = []
    for line in [l for l in text.split("\n") if l.strip()]:
        cs = line.split("；")
        for ci, cl in enumerate(cs):
            cl = cl.strip()
            if cl:
                out.append(cl + ("；" if ci < len(cs) - 1 else ""))
    return out

_SENT = re.compile(r"[^。；]*[。；]")
def split_body(text, target=SEG_MIN):
    """长段按句号/分号聚合成 ~target 字的短段。"""
    if len(text) <= target + 60:
        return [text]
    sents = _SENT.findall(text)
    tail = text[sum(len(s) for s in sents):]
    if tail.strip():
        sents.append(tail)
    out, cur = [], ""
    for s in sents:
        cur += s
        if len(cur) >= target:
            out.append(cur); cur = ""
    if cur.strip():
        if out and len(cur) < 40:
            out[-1] += cur
        else:
            out.append(cur)
    return out

def render_block(doc, kind, payload):
    if kind == "abstract":
        body_para(doc, norm_punct(payload))
    elif kind == "claim":
        for cl in split_claim(norm_punct(payload)):
            body_para(doc, cl)
    elif kind == "title":
        heading(doc, payload, "title")
    elif kind in ("h1", "h2"):
        heading(doc, payload, kind)
    elif kind == "body":
        for seg in split_body(norm_punct(payload)):
            body_para(doc, seg)
    elif kind == "code":
        add_code(doc, payload)

# ---------------- 分节 / 页眉 / 页脚 ----------------
def set_margins(section):
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.2); section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5); section.footer_distance = Cm(1.75)

def _clear_runs(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

def set_header(section, text):
    section.header.is_linked_to_previous = False
    p = section.header.paragraphs[0]; _clear_runs(p)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    for k, v in [("w:val", "single"), ("w:sz", "8"), ("w:space", "1"), ("w:color", "000000")]:
        bottom.set(qn(k), v)
    pBdr.append(bottom); pPr.insert(0, pBdr)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), size=15, bold=True)

def set_footer_pagenum(section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]; _clear_runs(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("- "), size=12)
    run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(it); run._r.append(e); set_font(run, size=12)
    set_font(p.add_run(" -"), size=12)

def clear_footer(section):
    section.footer.is_linked_to_previous = False
    _clear_runs(section.footer.paragraphs[0])

def restart_pgnum(section, start=1):
    sectPr = section._sectPr
    for el in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(el)
    pg = OxmlElement("w:pgNumType"); pg.set(qn("w:start"), str(start))
    # OOXML CT_SectPr 要求 pgNumType 在 cols/docGrid 之前;直接 append 会落到它们之后,
    # 生成非法序(MS Word 可能弹"内容无法读取/修复")。故插到 cols 前(无 cols 则 docGrid 前,再退 append)。
    cols = sectPr.find(qn("w:cols")); grid = sectPr.find(qn("w:docGrid"))
    (cols.addprevious(pg) if cols is not None
     else grid.addprevious(pg) if grid is not None else sectPr.append(pg))

# ---------------- 主流程 ----------------
def build(blocks_path, out_path):
    data = json.load(open(blocks_path, encoding="utf-8"))
    BLOCKS = data["blocks"]
    abs_blocks = [b for b in BLOCKS if b[0] == "abstract"]
    claim_blocks = [b for b in BLOCKS if b[0] == "claim"]
    fig_blocks = [b for b in BLOCKS if b[0] == "figure"]
    ti = next((i for i, b in enumerate(BLOCKS) if b[0] == "title"), None)
    spec_blocks = [b for b in (BLOCKS[ti:] if ti is not None else []) if b[0] != "figure"]

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = SONG; st.font.size = Pt(12)
    st.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), SONG)

    # 节1 说明书摘要
    s1 = doc.sections[0]
    set_margins(s1); set_header(s1, "说 明 书 摘 要"); clear_footer(s1)
    for k, pl in abs_blocks:
        render_block(doc, k, pl)

    # 节2 权利要求书
    s2 = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(s2); set_header(s2, "权 利 要 求 书"); set_footer_pagenum(s2); restart_pgnum(s2, 1)
    for k, pl in claim_blocks:
        render_block(doc, k, pl)

    # 节3 说明书
    s3 = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(s3); set_header(s3, "说 明 书"); set_footer_pagenum(s3); restart_pgnum(s3, 1)
    for k, pl in spec_blocks:
        render_block(doc, k, pl)

    # 节4 说明书附图
    s4 = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(s4); set_header(s4, "说 明 书 附 图"); clear_footer(s4)
    for idx, (k, pl) in enumerate(fig_blocks):
        if idx > 0:
            page_break(doc)
        path = pl if isinstance(pl, str) and not pl.lstrip().startswith("{") else json.loads(pl)["path"]
        add_figure(doc, path)

    doc.save(out_path)
    print("已生成:", out_path, "| 段落:", len(doc.paragraphs), "| 分节:", len(doc.sections))


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python build_patent_cnipa.py blocks.json 输出.docx"); sys.exit(1)
    build(sys.argv[1], sys.argv[2])
