# -*- coding: utf-8 -*-
"""inspect_template.py — 扒一份模板 docx 的真实版式,用来对齐格式(动笔/装配前先量)。

用法:
    python inspect_template.py 模板.docx

输出:分节数与分节符类型、各节页眉文字/页脚是否有页码与起始页、页面尺寸与页边距、
各类段落(标题/正文/权项)的字体字号对齐段距抽样、正文段长分布、媒体(图片)数。
据此核对 references/docx-format.md / cnipa-format-spec.md 的清单,并调整装配器参数。
**只读模板版式,不读取也不复制模板的技术内容。**
"""
import sys, re, zipfile, statistics
import docx
from docx.oxml.ns import qn


def main(path):
    z = zipfile.ZipFile(path)
    doc_xml = z.read("word/document.xml").decode("utf-8", "ignore")

    print("=" * 60)
    print("文件:", path)
    print("=" * 60)

    # 分节
    secs = re.findall(r"<w:sectPr.*?</w:sectPr>", doc_xml, re.S)
    print(f"\n[分节] 共 {len(secs)} 节")
    for i, s in enumerate(secs, 1):
        typ = re.findall(r'<w:type w:val="(\w+)"', s) or ["nextPage(默认)"]
        hdr = re.findall(r'headerReference[^>]*r:id="(\w+)"', s)
        ftr = re.findall(r'footerReference[^>]*r:id="(\w+)"', s)
        pg = re.findall(r'<w:pgNumType[^>]*w:start="(\d+)"', s)
        print(f"  节{i}: 分节符={typ} header={hdr} footer={ftr} 页码起始={pg or '无'}")

    # 页眉/页脚文字
    print("\n[页眉/页脚 文字]")
    for n in sorted(z.namelist()):
        if re.search(r"word/(header|footer)\d*\.xml$", n):
            x = z.read(n).decode("utf-8", "ignore")
            texts = "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", x))
            haspage = "PAGE" in x
            sz = re.findall(r'<w:sz w:val="(\d+)"', x)
            bdr = "pBdr" in x
            print(f"  {n}: 文字={texts!r} 含PAGE域={haspage} 字号(半磅)={sz[:1]} 下边框={bdr}")

    # 页面/页边距
    d = docx.Document(path); sec = d.sections[0]
    print("\n[页面] %.1f×%.1f cm  上%.1f 下%.1f 左%.1f 右%.1f  页眉距%.2f 页脚距%.2f (cm)" % (
        sec.page_width.cm, sec.page_height.cm, sec.top_margin.cm, sec.bottom_margin.cm,
        sec.left_margin.cm, sec.right_margin.cm,
        sec.header_distance.cm if sec.header_distance else -1,
        sec.footer_distance.cm if sec.footer_distance else -1))

    # 段落抽样:标题 vs 正文
    def ea_font(p):
        for r in p.runs:
            rPr = r.font.element.rPr
            if rPr is not None and rPr.rFonts is not None:
                v = rPr.rFonts.get(qn("w:eastAsia"))
                if v:
                    return v
        return None

    print("\n[段落抽样] (取若干非空段)")
    shown = 0
    for p in d.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        r0 = p.runs[0] if p.runs else None
        sz = r0.font.size.pt if r0 and r0.font.size else None
        bold = r0.font.bold if r0 else None
        pf = p.paragraph_format
        fi = pf.first_line_indent.pt if pf.first_line_indent else None
        print(f"  字号={sz} 粗={bold} 体={ea_font(p)} 对齐={p.alignment} 首行缩进={fi} ls={pf.line_spacing} | {t[:24]!r}")
        shown += 1
        if shown >= 14:
            break

    # 正文段长分布
    lens = [len(p.text) for p in d.paragraphs if len(p.text) > 40]
    if lens:
        print("\n[正文段长(>40字)] n=%d 最大=%d 中位=%d 均值=%.0f  最长3=%s" % (
            len(lens), max(lens), int(statistics.median(lens)), statistics.mean(lens),
            sorted(lens, reverse=True)[:3]))

    # 媒体(图片)数
    media = [n for n in z.namelist() if n.startswith("word/media/")]
    print("\n[媒体] 图片数:", len(media), "  (仅计数,不复制模板配图)")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python inspect_template.py 模板.docx"); sys.exit(1)
    main(sys.argv[1])
